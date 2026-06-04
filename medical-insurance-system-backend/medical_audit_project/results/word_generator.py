import os
from django.conf import settings
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn  # <-- 1. 把 qn 提到文件顶部
from docx.oxml import OxmlElement # <-- 1. 把 OxmlElement 也提到顶部
from django.http import HttpResponse
from datetime import datetime
from io import BytesIO
from urllib.parse import quote
from docx.shared import Pt, RGBColor, Cm 
from .models import Result

def set_font_style(run, font_name='宋体', size=None, bold=None, color_rgb=None, underline=None):
    """一个辅助函数，用于统一设置文字样式"""
    font = run.font
    font.name = font_name
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color_rgb:
        font.color.rgb = RGBColor.from_string(color_rgb)
    if underline is not None:
        font.underline = underline

def generate_task_report_docx(task_id, issue_number):
    """
    根据任务ID生成带特定样式的Word报告。
    """
    results = Result.objects.filter(task_id=task_id).select_related('rule').order_by('id')
    
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(3.6)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)
    # 设置默认中文字体和西文字体
    document.styles['Normal'].font.name = 'Times New Roman'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title_image_path = os.path.join(settings.BASE_DIR, 'results', 'static', 'results', 'images', 'title.png')
    underline_image_path = os.path.join(settings.BASE_DIR, 'results', 'static', 'results', 'images', 'underline.png')

    if os.path.exists(title_image_path):
        # 添加一个居中的段落来放置图片
        p_title_img = document.add_paragraph()
        p_title_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 插入图片，并设置宽度 (例如15厘米，你可以根据实际效果调整)
        p_title_img.add_run().add_picture(title_image_path, width=Cm(15))
    else:
        # 如果图片不存在，使用备用的文本标题
        title_p = document.add_paragraph('三医监管医疗行为数据简报')
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # ... (可以加上文本样式的代码)

    if os.path.exists(underline_image_path):
        p_underline_img = document.add_paragraph()
        p_underline_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_underline_img.add_run().add_picture(underline_image_path, width=Cm(15))
    # --- 2. 设置副标题 ---
    current_year = datetime.now().year
    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtitle_run = subtitle_p.add_run(f"【{current_year}年 第{issue_number}期】 内部资料")
    set_font_style(subtitle_run, font_name='仿宋_GB2312', size=16)

    document.add_paragraph() # 空一行

    # --- 3. 设置正文段落 ---
    p1 = document.add_paragraph("为强化医疗行为监管，保障医疗质量安全与医保基金合理使用，依托省三医联动平台，重点报告海南医学院第一附属医院的情况如下：")
    for run in p1.runs:
        set_font_style(run, font_name='仿宋_GB2312', size=16)

    # --- 4. 设置表格标题 ---
    document.add_paragraph()
    table_title_p = document.add_paragraph()
    table_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_title_run = table_title_p.add_run("表1  违规病历情况")
    set_font_style(table_title_run, font_name='黑体', size=12)

    # --- 5. 创建并设置表格 ---
    table_headers = ["序号", "医疗机构名称", "患者住院号", "违规类型", "违规项目", "违规原因", "有关依据"]
    table = document.add_table(rows=1, cols=len(table_headers))
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER # 表格整体居中
      # 总可用宽度约为 15.8 cm (A4 21cm - 左右边距 2.6*2)
    column_widths = {
        0: Cm(0.8),   # 序号 (窄)
        1: Cm(2.5),   # 医疗机构名称
        2: Cm(3.0),   # 患者住院号
        3: Cm(1.5),   # 违规类型
        4: Cm(1.5),   # 违规项目
        5: Cm(8.0),   # 违规原因 (显著加宽)
        6: Cm(2.5),   # 有关依据 (显著加宽)
    } # 总和: 0.8 + 2.5 + 3.0 + 2.0 + 2.5 + 2.5 + 2.5 = 15.8 cm
    # 设置表头样式
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(table_headers):
         # 设置列宽
        if i in column_widths:
            table.columns[i].width = column_widths[i]
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header_text)
        set_font_style(run, font_name='黑体', size=10.5, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4286f4') # 这是一个蓝色，你可以换成你想要的颜色代码
        tc_pr.append(shd)
    table.autofit = False
    table.layout_algorithm = 1 # 使用固定布局
    # 填充表格内容并设置样式
    for index, result in enumerate(results, start=1):
        row_cells = table.add_row().cells
        
        # 序号
        p = row_cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font_style(p.add_run(str(index)), font_name='宋体', size=10.5)
        
        # 其他列
        cell_data = [
            "海南医学院第一附属医院",
            result.hospitalization_id,
            result.rule.type if result.rule else "N/A",
            result.rule.drug_name if result.rule else "N/A",
            result.reason or "",
            result.rule.description if result.rule else "N/A",
        ]
        for i, text in enumerate(cell_data, start=1):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i <= 3 else WD_ALIGN_PARAGRAPH.LEFT # 前几列居中，后面居左
            set_font_style(p.add_run(text), font_name='宋体', size=10.5)

    # --- 6. 添加总结段落 ---
    document.add_paragraph()
    p2 = document.add_paragraph("结合以上数据：海南医学院第一附属医院的收费中存在如上疑似违规情况，建议结合临床实际开展调查，明确原因。")
    for run in p2.runs:
        set_font_style(run, font_name='仿宋_GB2312', size=16)

    # --- (保存到内存并返回 HttpResponse 的逻辑不变) ---
    from io import BytesIO
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)

    response = HttpResponse(
        doc_io.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    from urllib.parse import quote
    filename = f"违规报告_{current_year}_第{issue_number}期_任务{task_id}.docx"
    response['Content-Disposition'] = f"attachment; filename*=utf-8''{quote(filename)}"
    
    return response
from docx.oxml.ns import qn